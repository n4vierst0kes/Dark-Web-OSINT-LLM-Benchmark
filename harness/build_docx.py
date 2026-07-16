#!/usr/bin/env python3
"""Render the refined report to a .docx, filling Section IV from evaluation.json.

Abstract is reproduced verbatim (untouched). No Conclusion content is written
(reserved placeholder). Section IV uses the condensed single-table variant and
is populated with real per-configuration numbers when evaluation.json exists;
otherwise placeholders are emitted.
"""
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

HERE = Path(__file__).resolve().parent
EVAL = HERE / "evaluation.json"
OUT = HERE.parent / "DarkWeb_OSINT_Report.docx"

PRIMARY = "local::llama3.1:8b-instruct-q4_K_M"
ENTITIES = ["apt", "cves", "techniques", "goals", "country"]
SHORT = {  # display names for configs
    "local::llama3.1:8b-instruct-q4_K_M": "Llama 3.1 8B (Q4_K_M)",
    "local::mistral:7b-instruct-q4_K_M": "Mistral 7B (Q4_K_M)",
    "local::phi3:mini": "Phi-3 Mini (Q4)",
    "local::qwen2.5:1.5b-instruct-q4_K_M": "Qwen 2.5 1.5B (Q4_K_M)",
}
PH = "[pending]"


# ---------- helpers --------------------------------------------------------

def load_eval():
    if EVAL.exists():
        try:
            d = json.load(open(EVAL))
        except Exception:
            return None
        # only the new per-config schema is usable; older runs -> placeholders
        if isinstance(d, dict) and "configs" in d and d["configs"]:
            return d
    return None


def fmt(x, nd=3):
    if x is None:
        return "N/A"
    try:
        return f"{float(x):.{nd}f}"
    except Exception:
        return str(x)


def h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.italic = True
    r.font.size = Pt(10.5)
    return p


def body(doc, text, size=10, italic=False, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    return p


def eq(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.name = "Consolas"
    return p


def caption(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(9)
    return p


def make_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, htext in enumerate(headers):
        c = t.rows[0].cells[i]
        c.paragraphs[0].runs and c.paragraphs[0].runs[0].clear()
        run = c.paragraphs[0].add_run(htext)
        run.bold = True
        run.font.size = Pt(8.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(8.5)
    if widths:
        for i, w in enumerate(widths):
            for r in t.rows:
                r.cells[i].width = Inches(w)
    return t


# ---------- document -------------------------------------------------------

def build():
    ev = load_eval()
    doc = Document()
    # compact margins to keep page count down
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.7)
        s.left_margin = s.right_margin = Inches(0.7)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Dark Web OSINT: Evaluating Local and Cloud LLM Pipelines "
                  "for Dark Web Intelligence")
    r.bold = True
    r.font.size = Pt(16)

    authors = [
        ("Theofrolic Anathapindika Dean", "theofrolic.dean@binus.ac.id"),
        ("Felicia Wijaya", "felicia.wijaya010@binus.ac.id"),
        ("Calvin Martin", "calvin.martin@binus.ac.id"),
        ("Henry Lucky", "henry.lucky@binus.ac.id"),
        ("Galih Dea Pratama", "galih.pratama001@binus.ac.id"),
    ]
    ap = doc.add_paragraph()
    ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = ap.add_run(
        "  |  ".join(n for n, _ in authors)
        + "\nComputer Science Department, School of Computer Science, "
          "Bina Nusantara University, Jakarta, Indonesia 11480\n"
        + "  ".join(e for _, e in authors))
    ar.font.size = Pt(9)

    # Abstract — VERBATIM, UNTOUCHED
    pa = doc.add_paragraph()
    ra = pa.add_run("Abstract—")
    ra.bold = True
    ra.italic = True
    ra.font.size = Pt(9)
    ra2 = pa.add_run(
        "This electronic document is a “live” template and already "
        "defines the components of your paper [title, text, heads, etc.] in "
        "its style sheet. *CRITICAL: Do Not Use Symbols, Special Characters, "
        "Footnotes, or Math in Paper Title or Abstract. (Abstract)")
    ra2.italic = True
    ra2.font.size = Pt(9)

    pk = doc.add_paragraph()
    rk = pk.add_run("Keywords—")
    rk.bold = True
    rk.italic = True
    rk.font.size = Pt(9)
    pk.add_run(" Cyber threat intelligence (CTI), dark web, large language "
               "models, open-source intelligence (OSINT), performance "
               "benchmarking").font.size = Pt(9)

    # ---- SECTIONS (prose imported from the refined Markdown) ----
    for sec in SECTIONS:
        kind = sec[0]
        if kind == "h1":
            h1(doc, sec[1])
        elif kind == "h2":
            h2(doc, sec[1])
        elif kind == "p":
            body(doc, sec[1])
        elif kind == "eq":
            eq(doc, sec[1])
        elif kind == "cap":
            caption(doc, sec[1])
        elif kind == "tbl":
            make_table(doc, sec[1], sec[2], sec[3] if len(sec) > 3 else None)
        elif kind == "results":
            add_results_section(doc, ev)

    doc.save(OUT)
    print(f"wrote {OUT}")
    if ev is None:
        print("NOTE: evaluation.json not found -> Section IV emitted with placeholders")
    else:
        print("Section IV filled from evaluation.json")


def add_results_section(doc, ev):
    """Condensed Section IV: framing + single consolidated per-config table."""
    body(doc,
         "This section defines the evaluation framework and reporting schema "
         "for the local-versus-cloud comparison. Each corpus record is "
         "processed k = 3 times per backend at temperature 0.0 on the "
         "CPU-bound host of Section III-A. All local candidates (Table II) are "
         "evaluated under the simulated-hook mode over a static, manually "
         "annotated corpus in which every ground-truth label is verifiable in "
         "the input document. Because both backends receive byte-identical "
         "inputs, any difference is attributable to the model and its "
         "deployment boundary. Accuracy follows Eqs. (1)-(3) with Macro-F1 as "
         "the headline metric; performance, energy, cost, and exposure follow "
         "Eqs. (4)-(9); viability follows the four constraints of Section "
         "III-G with the Pareto frontier of Eq. (10). Values that hold by "
         "construction (ODES_local = 0, C_billed(local) = 0) are stated as "
         "definitional.")

    if ev is None:
        body(doc, "The benchmarking runs are still in progress; all "
                  "quantitative results below are pending.", italic=True)
        rows = [[SHORT[k], PH, PH, PH, PH, PH, "0*", PH]
                for k in SHORT]
        note = ("[Insert final benchmark data — per-entity accuracy, "
                "performance/cost/ODES, viability, Pareto set, and paired "
                "statistics — here once LLM evaluation completes.]")
        pareto_txt = "[pending]"
        stat_txt = "[pending]"
        per_entity_txt = "[pending]"
    else:
        cfg = ev["configs"]
        via = ev["viability"]
        rows = []
        for k in SHORT:
            if k not in cfg:
                rows.append([SHORT[k], "not run", "", "", "", "", "", ""])
                continue
            c = cfg[k]
            rows.append([
                SHORT[k],
                fmt(c["macro_f1"]),
                fmt(c["micro_f1"]),
                fmt(c["latency_p95_s"], 1),
                fmt(c["tps_mean"], 2),
                fmt(c["cost_total_usd"], 6),
                str(c["odes_bytes_total"]),
                "Yes" if via[k]["viable"] else "No",
            ])
        note = ("Cloud reference (OpenAI GPT-4o mini): not executed in this run "
                "— no billable provider key was available in the "
                "measurement environment; the harness is configured to "
                "populate the cloud column and the paired statistics once a "
                "key is supplied.")
        pf = ev.get("pareto", {})
        pareto_txt = ", ".join(SHORT.get(x, x) for x in pf.get("frontier", [])) or "n/a"
        stat_txt = ev.get("statistics", {}).get(
            "note", "paired local-vs-cloud test pending cloud execution")
        # per-entity for primary model
        pe = cfg.get(PRIMARY, {}).get("per_type", {})
        if pe:
            per_entity_txt = "; ".join(
                f"{e}: P={fmt(pe[e]['precision'],2)}, R={fmt(pe[e]['recall'],2)}, "
                f"F1={fmt(pe[e]['f1'],2)}" for e in ENTITIES)
        else:
            per_entity_txt = "n/a"

    caption(doc, "TABLE IV. CONSOLIDATED EVALUATION RESULTS (PER "
                 "CONFIGURATION). ODES in off-host bytes; 0 for all local "
                 "configurations by construction.")
    make_table(
        doc,
        ["Configuration", "Macro-F1", "Micro-F1", "p95 Lat (s)", "TPS",
         "C_total ($)", "ODES", "Viable"],
        rows,
        widths=[1.7, 0.7, 0.7, 0.8, 0.6, 0.9, 0.6, 0.6])

    body(doc, f"Per-entity accuracy (primary local model, {SHORT[PRIMARY]}): "
              f"{per_entity_txt}.", size=9)
    body(doc, f"Pareto-optimal configuration(s) (Eq. 10): {pareto_txt}.", size=9)
    body(doc, note, size=9, italic=True)

    h2(doc, "A. Statistical Analysis")
    body(doc,
         "Because the inference loop yields k ≥ 3 paired observations per "
         "record on both backends, the local-versus-cloud comparison is "
         "treated as a paired design. A paired t-test at α = 0.05 tests "
         "whether the mean per-record difference in Macro-F1 differs from "
         "zero, and Cohen's d is reported as the standardized effect size "
         "d = (x̄_local − x̄_cloud) / s_pooled. A non-significant "
         "test with negligible |d| indicates practical parity between the "
         "local SLM and the cloud API; a significant difference with a large "
         "effect size quantifies the capability gap a privacy-preserving "
         "deployment must accept.")
    body(doc, f"Paired-comparison result: {stat_txt}.", size=9, italic=True)


# ---------- static prose (condensed, refined) ------------------------------

SECTIONS = [
    ("h1", "I. INTRODUCTION"),
    ("p", "The modern cybersecurity landscape has made the Dark Web an "
     "essential source of Cyber Threat Intelligence (CTI) owing to its role "
     "as a primary hub for malware distribution, exploit trading, and the "
     "coordination of cyberattacks. The Onion Router (Tor) hidden services "
     "afford threat actors a high degree of privacy while complicating "
     "conventional intelligence-gathering. Large Language Models (LLMs) offer "
     "substantial automation potential in this domain, spanning threat "
     "detection, information extraction, and reasoning over complex security "
     "reports [1]. However, dark-web text exhibits fundamentally different "
     "vocabulary and structural characteristics, which limits the "
     "effectiveness of general-purpose models in producing accurate "
     "intelligence [2]. Open-Source Intelligence (OSINT) operations on the "
     "Dark Web are further constrained by an operational trilemma of legal "
     "liability, structural barriers from anti-crawling mechanisms, and "
     "limited practitioner resources [3]."),
    ("p", "Two deployment paradigms frame this problem. The first is a "
     "closed-source cloud pipeline, which offers high throughput but incurs "
     "operational-security risk: sensitive CTI data must leave the analyst's "
     "host, and adversarial prompts have been shown to extract memorized "
     "content [4], while precision degrades to approximately 0.76 on lengthy "
     "(~3,000-word) CTI reports [5]. The second is a small open-source "
     "pipeline, in which quantized 7B-8B-parameter models served locally "
     "impose no data-exposure risk but are constrained to 4-bit or 5-bit "
     "quantization within a 16 GB RAM environment without GPU acceleration "
     "[6]. The distinction is therefore one of model provenance and "
     "data-custody boundary, not physical hardware location."),
    ("p", "Prior surveys systematize attacks native to the Tor network and "
     "document the anonymization mechanisms that make attribution "
     "challenging [7], and catalogue analytical techniques for investigating "
     "dark-web content for CTI, highlighting a persistent gap between "
     "academic prototypes and production-grade pipelines [8]. At the "
     "multi-agent level, the MAD-CTI framework demonstrates autonomous "
     "scraping, classification, and summarization of dark-web content, "
     "improving consistency over single-model pipelines [9]."),
    ("p", "Throughout this paper, local pipeline and cloud pipeline refer to "
     "model provenance and data-custody boundaries, not physical hardware "
     "location. The local pipeline deploys small open-source quantized models "
     "(Llama 3.1 8B, Mistral 7B, Phi-3 Mini) on analyst-controlled hardware "
     "via Ollama, so all computation and data remain on-device. The cloud "
     "pipeline routes inference through a closed-source proprietary API "
     "(OpenAI GPT-4o mini), where prompt data crosses a network boundary to a "
     "third party. This framing follows [6]: the operative difference is "
     "whether sensitive CTI data leaves the analyst's custody."),
    ("p", "The present study pursues four objectives. First, to operationalize "
     "a reproducible benchmarking harness that routes identical "
     "CTI-extraction prompts through both a locally served quantized SLM and "
     "a cloud API backend, using a static, pre-annotated ground-truth corpus "
     "in place of live scraping. Second, to measure the capability gap across "
     "Precision, Recall, F1, latency, throughput, and energy-per-task on "
     "representative analyst hardware. Third, to quantify the "
     "operational-security cost of the cloud pipeline via a defined OpSec "
     "Data Exposure Score (ODES). Fourth, to derive a composite "
     "cost-performance-accuracy analysis, including a Pareto-frontier "
     "computation and three composite metrics (F1-per-dollar, "
     "cost-per-correct-extraction, and cost-adjusted viability), enabling an "
     "analyst to select a configuration under an explicit, recorded "
     "preference rule. The outcome is a measurement procedure whose output is "
     "a workload-specific, hardware-specific, preference-explicit answer to "
     "the local-versus-cloud decision."),

    ("h1", "II. RELATED WORK"),
    ("p", "Dark-web OSINT faces the trilemma of legal concerns, Tor anonymity, "
     "and limited resources [3]. Dark-web language diverges systematically in "
     "vocabulary: DarkBERT shows domain-specific fine-tuning yields "
     "improvements of up to 13 F1 points on downstream classification [2], "
     "and the CoDA dataset provides a 10k-document corpus for studying "
     "dark-web language [10]. MAD-CTI demonstrates autonomous multi-agent "
     "scraping, classification, and summarization [9], while zero-shot "
     "classification with commercial LLMs is competitive on illicit-content "
     "categorization yet disagrees widely across models [11]."),
    ("p", "For structured CTI extraction, the Vulcan BERT+BiLSTM framework "
     "attains an NER F1 of 0.972 on labeled security articles [12]. LLM-based "
     "filtering of MITRE ATT&CK technique candidates reduces the candidate "
     "set by roughly 33x while retaining up to 94% of ground-truth labels "
     "[13], and prompt-based methods enable structured threat-relationship "
     "extraction from unstructured CTI [14]. Trust in LLMs for cybersecurity "
     "is constrained by hallucination on lengthy documents; precision "
     "degrades to approximately 0.76 on ~3,000-word CTI reports [5]. "
     "Accordingly, contemporary practice adopts Macro-F1 as a standard "
     "accuracy measure, which this study employs to compare output stability "
     "between cloud and local models. Data security is co-equal, as cloud "
     "models are vulnerable to extraction of sensitive information via "
     "adversarial prompts [15]; the literature advocates data isolation "
     "through local deployment, leaving open whether privacy constraints "
     "necessarily degrade intelligence quality."),
    ("p", "Beyond entity extraction, structured attack graphs can be built "
     "from large-scale CTI reports [16] and further automated through "
     "zero-shot LLM guidance [17]. Cloud-API risks such as prompt injection "
     "and PII leakage are well-documented structural threats [22] and "
     "motivate the ODES measure [24]. Systematic surveys indicate the "
     "capability gap between open-source and proprietary models remains "
     "unresolved for cybersecurity tasks [20], reinforced across more than "
     "ten downstream applications [22]. Class-aware rebalancing is effective "
     "for ATT&CK technique classification and serves as a benchmark reference "
     "[21]. The maturity gap between frontier and open-weight models has been "
     "assessed for automated cybersecurity tasks [23] and verified "
     "empirically for CVE mapping, ATT&CK technique extraction, and "
     "threat-actor attribution [30]. Deploying LLMs on consumer hardware "
     "demands aggressive optimization: within a 16 GB, GPU-less environment, "
     "4-bit quantization (e.g., Q4_K_M) is a technical prerequisite for "
     "7B/8B-class models [6]. The benchmarking methodology here follows the "
     "CPU-bound edge-LLM protocol of [26], complemented by the LLM-based "
     "APT-detection taxonomy of [25]."),

    ("h1", "III. RESEARCH METHODOLOGY"),
    ("p", "This section describes the systematic methodology used to evaluate "
     "and compare local and cloud LLM pipelines for Dark Web OSINT. The "
     "framework is organized into six sequential stages that isolate the "
     "trade-offs among extraction accuracy, system latency, execution cost, "
     "and operational-security exposure. The orchestration layer (the Robin "
     "Framework) coordinates data ingestion, model execution, and telemetry "
     "collection across both backends. By standardizing prompt templates and "
     "input corpora, any variance in accuracy or performance is attributable "
     "to the model architecture and deployment boundary rather than to "
     "engineering discrepancies. Table I summarizes the stages."),
    ("cap", "TABLE I. EVALUATION METHODOLOGY: SIX-STAGE PIPELINE"),
    ("tbl",
     ["Stage", "Concern", "Principal Output"],
     [["1 - Environment & Orchestration",
       "CPU-bound host; Tor path decoupled from inference",
       "Controlled, reproducible runtime"],
      ["2 - Model Configuration",
       "Local Q4_K_M GGUF via Ollama; cloud REST API",
       "Backend adapters; custody boundary"],
      ["3 - Script Architecture (simulated-hook)",
       "Live Tor scraping stubbed; static corpus inputs",
       "Byte-identical inputs; LLM as sole variable"],
      ["4 - Inference Loop & Determinism",
       "Temperature 0.0; k >= 3 repetitions per record",
       "Residual non-determinism probe; paired samples"],
      ["5 - Instrumentation & Timing",
       "Per-phase perf_counter_ns hooks",
       "Load, generation, and parse latencies"],
      ["6 - Scoring & Viability",
       "Accuracy, performance, ODES, cost; threshold check",
       "Per-configuration viability and Pareto set"]],
     [1.7, 2.4, 2.4]),

    ("h2", "A. Experimental Setup and Infrastructure"),
    ("p", "Local inference runs on a single analyst-grade host - a laptop with "
     "16 GB DDR4 RAM, an x86-64 CPU, and no discrete GPU - matching the "
     "consumer-hardware class benchmarked in [26]. All local inference is "
     "therefore CPU-bound. Network egress for any Tor-facing component is "
     "routed through a SOCKS5 proxy bound to 127.0.0.1:9050; this path is not "
     "exercised during measurement, so reported timings reflect inference "
     "behaviour alone rather than variable Tor latency. The open-source "
     "pipeline is served via Ollama, and the closed-source pipeline calls the "
     "OpenAI REST API; both consume the same prompt templates to eliminate "
     "engineering-induced variance [27]."),
    ("h2", "B. Corpus and Dataset Infrastructure"),
    ("p", "The evaluation dataset is the Di Tizio APT Bundle, available on "
     "Zenodo (DOI: 10.5281/zenodo.6514817) [28], containing genuine threat "
     "reports in PDF form together with structured ground-truth labels "
     "(threat actors, CVEs, MITRE ATT&CK techniques, aliases, country). A "
     "static, pre-annotated corpus is used in place of live dark-web scraping "
     "to guarantee reproducibility and avoid Tor-induced latency and content "
     "drift [3]. The ground truth is text-grounded: each label is verifiable "
     "in the input document shown to the model, so that recall reflects "
     "extraction capability rather than corpus mismatch. Fixing the corpus "
     "ensures repeated executions are scored against byte-identical "
     "references, the precondition for reproducibility."),
    ("h2", "C. Model Selection and Resource Tuning"),
    ("p", "Table II lists the model candidates. Llama 3.1 8B (Q4_K_M) is the "
     "primary local reference; Mistral 7B, Phi-3 Mini, and a compact "
     "1.5B-class model are secondary candidates spanning the "
     "memory-capability trade-off. OpenAI GPT-4o mini is the closed-source "
     "reference. All local models use 4-bit quantization as the operative "
     "lever for fitting 7B/8B-class weights within the 16 GB envelope on a "
     "GPU-less host [6]."),
    ("cap", "TABLE II. LOCAL AND CLOUD MODEL CANDIDATES AND MEMORY FOOTPRINT"),
    ("tbl",
     ["Model Tag", "Quantization", "Approx. RAM", "Throughput", "Status"],
     [["llama3.1:8b-instruct-q4_K_M", "Q4_K_M (4-bit)", "~5.5 GB",
       "8-15 t/s [26]", "Primary local reference"],
      ["mistral:7b-instruct-q4_K_M", "Q4_K_M (4-bit)", "~4.1 GB",
       "10-18 t/s [26]", "Secondary candidate"],
      ["phi3:mini", "Q4 (4-bit)", "~2.0 GB", "18-30 t/s [26]",
       "Secondary candidate"],
      ["qwen2.5:1.5b-instruct-q4_K_M", "Q4_K_M (4-bit)", "~1.0 GB",
       "30-60 t/s [26]", "Low-footprint SLM"],
      ["gpt-4o-mini (cloud)", "Provider-managed", "N/A",
       "60-90 t/s", "Cloud reference"]],
     [2.3, 1.3, 0.9, 1.0, 1.5]),
    ("p", "The candidates reflect common OSINT deployment options under the "
     "stated hardware constraints. Llama 3.1 8B is the primary local model "
     "given its reasoning capability under 4-bit quantization on commodity "
     "hardware. Mistral 7B balances accuracy, speed, and memory footprint, "
     "enabling a within-tier local comparison. Phi-3 Mini and the 1.5B-class "
     "model represent the low-footprint SLM regime for the most constrained "
     "hosts. OpenAI GPT-4o mini is the cloud reference for its low latency and "
     "large context window."),
    ("h2", "D. CTI Entity Extraction Protocol"),
    ("p", "A single fixed prompt template is applied across all backends [27], "
     "requesting five fields in canonical form: apt (APT-group names), cves, "
     "techniques (MITRE ATT&CK canonical names), goals, and country. Decoding "
     "temperature is fixed at 0.0 (greedy decoding), and each record is "
     "processed k = 3 times per backend. Because temperature is fixed, the "
     "repetitions probe residual non-determinism - backend scheduling, "
     "tokenizer edge cases, and provider-side variability - directly "
     "addressing the instability reported by [5]. For each entity type, true "
     "positives are identified by case-insensitive set intersection between "
     "the predicted and ground-truth lists. For the apt type, alias expansion "
     "merges all known aliases into the reference before comparison: "
     "extracting Fancy Bear, Sofacy, or STRONTIUM all count as true positives "
     "when the ground truth records APT28, reflecting operational CTI "
     "practice."),
    ("h2", "E. Benchmarking Metrics: Extraction Accuracy"),
    ("p", "Extraction accuracy is quantified using Precision, Recall, and F1, "
     "computed per entity type from the parsed output (recovered by a "
     "regular-expression and JSON parsing layer) against the ground-truth "
     "labels, following the automated TP/FP/FN accounting of [12]:"),
    ("eq", "Precision = TP / (TP + FP)                                  (1)"),
    ("eq", "Recall    = TP / (TP + FN)                                  (2)"),
    ("eq", "F1 = 2 . (Precision . Recall) / (Precision + Recall)        (3)"),
    ("p", "Metrics are computed per entity type and aggregated. Macro-F1 - the "
     "unweighted mean of per-type F1 - is the primary headline metric, "
     "because entity frequencies in CTI corpora are highly imbalanced. "
     "Micro-F1 (over pooled TP/FP/FN) is a secondary metric. Both are "
     "established in [5], [29]."),
    ("h2", "F. Benchmarking Metrics: Performance, Energy, Cost, and ODES"),
    ("p", "Timing is captured in-process with the monotonic, "
     "nanosecond-resolution clock perf_counter_ns, with hooks bracketing "
     "three phases of every call: (i) model load / connection, (ii) query and "
     "generation, and (iii) JSON parsing, following [26]. For each call i, "
     "throughput is:"),
    ("eq", "throughput_i = output_tokens_i / latency_i                  (4)"),
    ("p", "Mean and 95th-percentile (p95) latency are reported; p95 governs "
     "analyst responsiveness [26]. The operational target for the local "
     "pipeline is an end-to-end latency below 180 s per record. Local-pipeline "
     "energy is measured via Intel RAPL, with a wall-clock x TDP fallback "
     "when RAPL is unavailable; cloud energy is provider-absorbed and is not "
     "compared directly [30]. Per-task cost follows the Green-AI principle "
     "that financial cost is a first-class criterion [31]:"),
    ("eq", "C_energy = (E_joules / 3,600,000) . P_kWh                   (5)"),
    ("eq", "C_billed = (in_tok/1e6).R_in + (out_tok/1e6).R_out          (6)"),
    ("eq", "F1_per_dollar = macro_F1 / C_total                          (7)"),
    ("p", "For the local pipeline C_billed = 0; for the cloud pipeline "
     "C_energy = 0; C_total = C_energy + C_billed, and C_per_TP = C_total / "
     "TP_count. The OpSec Data Exposure Score (ODES) quantifies "
     "investigation-relevant data leaving the host; under the "
     "adversarial-extraction threat model of [18], off-host volume is the "
     "conservative exposure proxy:"),
    ("eq", "ODES_local = 0                                              (8)"),
    ("eq", "ODES_cloud = S . B_cloud                                    (9)"),
    ("p", "where B_cloud is the total bytes of prompt content transmitted to "
     "the cloud endpoint and S is a unitless per-byte sensitivity coefficient "
     "(default S = 1, reducing ODES to raw off-host bytes). Table III "
     "summarizes all metrics."),
    ("cap", "TABLE III. SUMMARY OF BENCHMARKING METRICS"),
    ("tbl",
     ["Metric", "Formula / Anchor", "Optimum"],
     [["Macro-F1", "Mean of per-type F1 [5],[10]", "Higher"],
      ["Micro-F1", "F1 over pooled TP/FP/FN [1],[10]", "Higher"],
      ["Mean / p95 latency", "Wall-clock s per call [26]", "Lower"],
      ["Throughput", "output_tokens / latency (TPS) [26]", "Higher"],
      ["Energy (J)", "dRAPL or TDP x t [26]", "Lower"],
      ["ODES", "S . B_cloud off-host bytes [4],[18]", "Lower (0 local)"],
      ["C_total (USD)", "C_energy + C_billed [9],[31]", "Lower"],
      ["F1/$ ; C_per_TP", "macro_F1/C_total ; C_total/TP [9]", "Higher; Lower"]],
     [1.7, 3.1, 1.4]),
    ("h2", "G. Operational Viability Threshold"),
    ("p", "A pipeline is operationally viable for a task if and only if all "
     "four analyst-defined constraints are simultaneously satisfied: macro-F1 "
     ">= F1_min [5],[12]; p95 latency <= L_max (180 s) [26]; ODES <= ODES_max "
     "[18],[32]; and C_total <= C_max [31]. Beyond binary viability, the "
     "harness computes the Pareto frontier across all (model, quantization, "
     "backend) configurations on accuracy A = macro-F1 (higher better), "
     "exposure O = ODES (lower better), and cost L = C_total (lower better). "
     "A configuration m' Pareto-dominates m iff:"),
    ("eq", "A_m' >= A_m  AND  O_m' <= O_m  AND  L_m' <= L_m            (10)"),
    ("p", "with at least one strict inequality. Pareto-dominated "
     "configurations are eliminated without loss [6], operationalizing the "
     "[31] principle that no single scalar metric should govern selection."),

    ("h1", "IV. RESULTS AND DISCUSSION"),
    ("results",),

    ("h1", "V. CONCLUSION"),
    ("p", "[Reserved - to be completed once final evaluation results are "
     "available.]"),

    ("h1", "VI. AUTHOR STATEMENT"),
    ("p", "Theofrolic Anathapindika Dean contributed to conceptualization, "
     "methodology, software development, formal analysis, writing of the "
     "original draft, and review and editing. Felicia Wijaya contributed to "
     "data curation, validation, and review and editing. Calvin Martin "
     "contributed to visualization and review and editing. Henry Lucky "
     "contributed to supervision and review and editing. Galih Dea Pratama "
     "contributed to supervision and review and editing."),

    ("h1", "REFERENCES"),
]

REFS = [
    "[1] Y. Chen et al., “A survey of large language models for cyber threat detection,” 2024.",
    "[2] Y. Jin et al., “DarkBERT: A Language Model for the Dark Side of the Internet,” Proc. ACL, vol. 1, pp. 7515-7533, 2023, doi: 10.18653/v1/2023.acl-long.415.",
    "[3] P. Kuhn, K. Wittorf, and C. Reuter, “Navigating the Shadows,” IEEE Access, vol. 12, pp. 118903-118922, 2024, doi: 10.1109/ACCESS.2024.3448247.",
    "[4] B. C. Das, M. H. Amini, and Y. Wu, “Security and Privacy Challenges of Large Language Models: A Survey,” ACM Comput. Surv., vol. 57, no. 6, pp. 1-39, 2025, doi: 10.1145/3712001.",
    "[5] E. Mezzi, F. Massacci, and K. Tuma, Large Language Models Are Unreliable for Cyber Threat Intelligence, vol. 15993 LNCS. Springer, 2025, doi: 10.1007/978-3-032-00627-1_17.",
    "[6] Y. Zheng et al., “A Review on Edge Large Language Models,” ACM Comput. Surv., vol. 57, no. 8, 2024, doi: 10.1145/3719664.",
    "[7] J. Saleem, R. Islam, and M. A. Kabir, “The Anonymity of the Dark Web: A Survey,” IEEE Access, vol. 10, pp. 33628-33660, 2022, doi: 10.1109/ACCESS.2022.3161547.",
    "[8] R. Basheer and B. Alkhatib, “Threats from the Dark,” J. Comput. Netw. Commun., vol. 2021, pp. 1-21, 2021, doi: 10.1155/2021/1302999.",
    "[9] S. Shah and V. K. Madisetti, “MAD-CTI,” IEEE Access, vol. 13, pp. 40158-40168, 2025, doi: 10.1109/ACCESS.2025.3547172.",
    "[10] Y. Jin et al., “Shedding New Light on the Language of the Dark Web,” Proc. NAACL-HLT, pp. 5621-5637, 2022, doi: 10.18653/v1/2022.naacl-main.412.",
    "[11] V.-P. Prado-Sánchez et al., “Zero-Shot Classification of Illicit Dark Web Content with Commercial LLMs,” Electronics, vol. 14, no. 20, p. 4101, 2025, doi: 10.3390/electronics14204101.",
    "[12] H. Jo, Y. Lee, and S. Shin, “Vulcan,” Comput. Secur., vol. 120, p. 102763, 2022, doi: 10.1016/j.cose.2022.102763.",
    "[13] A. Krašovec et al., “Large Language Models for Cyber Threat Intelligence: Extracting MITRE With LLMs,” pp. 80-89, 2025, doi: 10.1007/978-3-032-00633-2_5.",
    "[14] X. Liu and Z. Ding, “CyberRE-LLM,” pp. 261-271, 2025, doi: 10.1007/978-981-96-9994-0_22.",
    "[15] A. Colacicco et al., “Exploring Approaches for Detecting Memorization of Recommender System Data in LLMs,” 2026.",
    "[16] Z. Wang et al., “Automated Attack Knowledge Graph Construction with LLMs,” Proc. 2025 2nd Int. Conf. Computer and Multimedia Technology, pp. 700-706, 2025, doi: 10.1145/3757749.3757864.",
    "[17] Z. Li, J. Zeng, Y. Chen, and Z. Liang, “AttacKG,” pp. 589-609, 2022, doi: 10.1007/978-3-031-17140-6_29.",
    "[18] A. Kucharavy et al., Large Language Models in Cybersecurity. Springer, 2024, doi: 10.1007/978-3-031-54827-7.",
    "[19] K. Edemacu and X. Wu, “Privacy Preserving Prompt Engineering: A Survey,” ACM Comput. Surv., vol. 57, no. 10, pp. 1-36, 2025, doi: 10.1145/3729219.",
    "[20] H. Xu et al., “Large Language Models for Cyber Security: A Systematic Literature Review,” ACM Trans. Softw. Eng. Methodol., 2025, doi: 10.1145/3769676.",
    "[21] M. Albarrak, A. Alqudhaibi, and S. Jagtap, “AC_MAPPER,” Int. J. Inf. Secur., vol. 24, no. 6, p. 232, 2025, doi: 10.1007/s10207-025-01146-5.",
    "[22] J. Zhang et al., “When LLMs meet cybersecurity,” Cybersecurity, vol. 8, no. 1, p. 55, 2025, doi: 10.1186/s42400-025-00361-w.",
    "[23] T. Conceição and N. Cruz, “Evaluation of the maturity of LLMs in the cybersecurity domain,” Int. J. Inf. Secur., vol. 24, no. 5, p. 197, 2025, doi: 10.1007/s10207-025-01112-1.",
    "[24] M. T. Alam et al., “CTIBench,” Adv. Neural Inf. Process. Syst. 37, pp. 50805-50825, 2024, doi: 10.52202/079017-1607.",
    "[25] M. Golec et al., “LLM-Driven APT Detection for 6G Wireless Networks,” IEEE Access, vol. 13, pp. 145271-145288, 2025, doi: 10.1109/ACCESS.2025.3595665.",
    "[26] E. J. Husom et al., “Sustainable LLM Inference for Edge AI,” ACM Trans. Internet Things, vol. 6, no. 4, 2025, doi: 10.1145/3767742.",
    "[27] N. Tihanyi et al., “CyberMetric,” Proc. IEEE CSR 2024, pp. 296-302, 2024, doi: 10.1109/CSR61664.2024.10679494.",
    "[28] G. D. Tizio, M. Armellini, and F. Massacci, “Software Updates Strategies,” IEEE Trans. Softw. Eng., vol. 49, no. 3, pp. 1359-1373, 2023, doi: 10.1109/TSE.2022.3176674.",
    "[29] Y. Chen et al., “A survey of large language models for cyber threat detection,” Comput. Secur., vol. 145, p. 104016, 2024, doi: 10.1016/j.cose.2024.104016.",
    "[30] D. Patterson et al., “Energy and Emissions of Machine Learning on Smartphones vs. the Cloud,” Commun. ACM, vol. 67, no. 2, pp. 86-97, 2024, doi: 10.1145/3624719.",
    "[31] R. Schwartz, J. Dodge, N. A. Smith, and O. Etzioni, “Green AI,” Commun. ACM, vol. 63, no. 12, pp. 54-63, 2020, doi: 10.1145/3381831.",
    "[32] A. Colacicco et al., “Exploring Approaches for Detecting Memorization of Recommender System Data in LLMs,” 2025.",
]


def add_refs(doc):
    for r in REFS:
        p = doc.add_paragraph()
        run = p.add_run(r)
        run.font.size = Pt(8.5)


# append refs after building main sections
_orig_build = build


def build_with_refs():
    ev = load_eval()
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.7)
        s.left_margin = s.right_margin = Inches(0.7)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Dark Web OSINT: Evaluating Local and Cloud LLM Pipelines "
                  "for Dark Web Intelligence")
    r.bold = True; r.font.size = Pt(16)
    authors = ["Theofrolic Anathapindika Dean", "Felicia Wijaya",
               "Calvin Martin", "Henry Lucky", "Galih Dea Pratama"]
    emails = ["theofrolic.dean", "felicia.wijaya010", "calvin.martin",
              "henry.lucky", "galih.pratama001"]
    ap = doc.add_paragraph(); ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = ap.add_run("  |  ".join(authors) + "\nComputer Science Department, "
                    "School of Computer Science, Bina Nusantara University, "
                    "Jakarta, Indonesia 11480\n"
                    + "  ".join(e + "@binus.ac.id" for e in emails))
    ar.font.size = Pt(9)

    pa = doc.add_paragraph()
    ra = pa.add_run("Abstract—"); ra.bold = True; ra.italic = True; ra.font.size = Pt(9)
    ra2 = pa.add_run(
        "This electronic document is a “live” template and already "
        "defines the components of your paper [title, text, heads, etc.] in "
        "its style sheet. *CRITICAL: Do Not Use Symbols, Special Characters, "
        "Footnotes, or Math in Paper Title or Abstract. (Abstract)")
    ra2.italic = True; ra2.font.size = Pt(9)
    pk = doc.add_paragraph()
    rk = pk.add_run("Keywords—"); rk.bold = True; rk.italic = True; rk.font.size = Pt(9)
    pk.add_run(" Cyber threat intelligence (CTI), dark web, large language "
               "models, open-source intelligence (OSINT), performance "
               "benchmarking").font.size = Pt(9)

    for sec in SECTIONS:
        kind = sec[0]
        if kind == "h1":
            h1(doc, sec[1])
        elif kind == "h2":
            h2(doc, sec[1])
        elif kind == "p":
            body(doc, sec[1])
        elif kind == "eq":
            eq(doc, sec[1])
        elif kind == "cap":
            caption(doc, sec[1])
        elif kind == "tbl":
            make_table(doc, sec[1], sec[2], sec[3] if len(sec) > 3 else None)
        elif kind == "results":
            add_results_section(doc, ev)
    add_refs(doc)

    doc.save(OUT)
    print(f"wrote {OUT}")
    print("Section IV: " + ("filled from evaluation.json" if ev else "placeholders"))


if __name__ == "__main__":
    build_with_refs()
